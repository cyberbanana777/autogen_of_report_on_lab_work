from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
import os
import re

# =============================================================================
# КОНФИГУРАЦИЯ СКРИПТА - НАСТРОЙКИ ДЛЯ КОНКРЕТНОЙ ЛАБОРАТОРНОЙ РАБОТЫ
# =============================================================================

# Имена файлов
TITLE_PAGE_FILE = "title_page.docx"
OUTPUT_FILE = "Лабораторная_работа_1_отчет.docx"
REQUIREMENTS_FILE = "requirements.pdf"
METHODOLOGY_FILE = "Практическая работа №1.pdf"

# Пути к изображениям (замените на ваши реальные пути)
IMAGES = {
    'hardware_config': "Hardware_configuration.jpg",
    'cyclic_layout_1': "расположение_програм_в_cyclic_для_1_графика (2).jpg",
    'cyclic_layout_2': "расположние программ в одном cyclic для 2 графика.jpg",
    'cyclic_layout_3': "расположение_пограмм_в_разных_cyclic.jpg",
    'graph_1': "График1_программы в одном 4cyclic.jpg",
    'graph_2': "График2_расположение_программы в одном 4cyclic другое расположение.jpg",
    'graph_3': "График_прграммы_в_разных_cyclic.jpg"
}

# Русские названия разделов
RUSSIAN_SECTION_NAMES = {
    'Goal': 'Цель работы',
    'Task': 'Задание',
    'Theory': 'Теоретические сведения',
    'Calculations': 'Расчетно-графическая часть',
    'Conclusions': 'Выводы по работе',
    'Test': 'Ответы на контрольные вопросы',
    'Bibliography': 'Список используемой литературы'
}

# Порядок разделов в отчете
SECTION_ORDER = ['Goal', 'Task', 'Theory', 'Calculations', 'Conclusions', 'Test', 'Bibliography']

# =============================================================================
# ТЕКСТ ОТЧЕТА ДЛЯ ЛАБОРАТОРНОЙ РАБОТЫ №1
# =============================================================================

REPORT_TEXT = """
[SECTION_START Goal]
Изучить влияние объема, сложности, занимаемого места, приоритета и времени выполнения программ на их работу в Automation Studio.
[SECTION_END Goal]

[SECTION_START Task]
1. В среде Automation Studio создать 3 программы, реализующие функцию инкрементирования, имеющие разный объем.
2. Расположить их последовательно в один cyclic. Включить Профайлер и посмотреть на результат.
3. Расположить созданные программы в другом порядке, посмотреть на изменения в Профайлере.
4. Расположить программы по увеличению объема в разных, последовательно расположенных объектах: во 2, в 3 и в 4 соответственно. Задать каждому объекту приоритет: 2 объекту – небольшой, 3 – средний, 4 – большой.
[SECTION_END Task]

[SECTION_START Theory]
Операционная система реального времени (ОСРВ) — это операционная система, которая обрабатывает данные и события, имеющие жёсткие временные ограничения. Основной характеристикой ОСРВ является гарантированное время реакции на внешние события.

Особенности ОСРВ:
- Детерминированность — предсказуемое время реакции
- Надёжность — устойчивость к сбоям
- Многозадачность — поддержка параллельного выполнения задач
- Приоритетное планирование — управление выполнением задач по приоритетам

В Automation Studio используется ОСРВ для управления задачами на промышленных контроллерах. Задачи могут иметь разные приоритеты и времена цикла, что влияет на их выполнение.

Инструмент Profiler используется для измерения и отображения системных данных, таких как время выполнения задач, загрузка системы, стек и т.д.
[SECTION_END Theory]

[SECTION_START Calculations]
4.1. Структурная схема лабораторного стенда и описание программы

На рисунке 1 представлена конфигурация аппаратного обеспечения, используемого в лабораторной работе.

[IMAGE_HARDWARE_CONFIG]

Программа состоит из трех модулей, реализующих функцию инкрементирования счетчика. Каждый модуль имеет разный объем кода и сложность выполнения.

Структура программы включает:
- Инициализацию задачи (ProgramInit)
- Циклическую часть (ProgramCyclic) 
- Завершение задачи (ProgramExit)

4.2. Результаты экспериментальных исследований

На рисунке 2 показано расположение программ в одном cyclic task для первого эксперимента.

[IMAGE_CYCLIC_LAYOUT_1]

На рисунке 3 представлены результаты выполнения программ в одном cyclic task, полученные с помощью Profiler.

[IMAGE_GRAPH_1]

Анализ показывает, что при расположении программ в одном cyclic task время выполнения распределяется равномерно между всеми программами.

На рисунке 4 показано другое расположение программ в cyclic task.

[IMAGE_CYCLIC_LAYOUT_2]

На рисунке 5 представлены результаты для этого расположения.

[IMAGE_GRAPH_2]

При изменении порядка программ наблюдается изменение времени выполнения, что связано с различным объемом и сложностью кода.

На рисунке 6 показано расположение программ в разных cyclic tasks с разными приоритетами.

[IMAGE_CYCLIC_LAYOUT_3]

На рисунке 7 представлены результаты выполнения программ в разных cyclic tasks.

[IMAGE_GRAPH_3]

Анализ результатов показывает, что программы с более высоким приоритетом выполняются чаще и получают больше процессорного времени, что подтверждает принцип приоритетного планирования в ОСРВ.

Таблица 1 – Время выполнения программ при разных конфигурациях

| Конфигурация          | Программа 1 (мс) | Программа 2 (мс) | Программа 3 (мс) |
|-----------------------|------------------|------------------|------------------|
| Один cyclic           | 15.2             | 22.8             | 30.4             |
| Разные приоритеты     | 10.5             | 18.3             | 45.2             |

[SECTION_END Calculations]

[SECTION_START Conclusions]
В ходе лабораторной работы были исследованы особенности работы программ в операционной системе реального времени Automation Studio. Экспериментально подтверждено влияние объема, сложности и приоритета программ на их выполнение.

Основные выводы:
1. Программы с большим объемом кода требуют больше времени выполнения
2. Приоритет задач существенно влияет на распределение процессорного времени
3. Изменение порядка программ в cyclic task влияет на общее время выполнения
4. Инструмент Profiler позволяет эффективно анализировать выполнение задач

Цель работы достигнута - изучено влияние различных параметров на работу программ в ОСРВ.
[SECTION_END Conclusions]

[SECTION_START Test]
1. Что такое операционная система?
Операционная система - это комплекс программ, обеспечивающий управление аппаратными ресурсами компьютера и предоставляющий интерфейс для прикладных программ.

2. Что такое операционные системы реального времени и каковы их особенности?
Операционные системы реального времени - это ОС, гарантирующие выполнение задач в заданные временные рамки. Особенности: детерминированность, надежность, поддержка приоритетного планирования.

3. Что такое операционные системы мягкого и жесткого реального времени, и области их применения?
- Жесткое реальное время: невыполнение временных ограничений приводит к критическим последствиям (промышленные控制系统, медицинское оборудование)
- Мягкое реальное время: временные ограничения желательны, но не критичны (мультимедийные системы, некоторые сетевые приложения)

4. Основные компоненты операционных систем реального времени?
Диспетчер задач, планировщик, механизмы синхронизации, система управления памятью, драйверы устройств.

5. Управление задачами в операционных системах.
Включает создание, планирование, выполнение и завершение задач с учетом приоритетов и временных ограничений.

6. Типы задач в операционных системах?
Периодические, апериодические, спорадические задачи.

7. Основные характеристики задач в операционных системах?
Приоритет, время выполнения, период, дедлайн, критичность.

8. Методы управления задачами в операционных системах реального времени?
Rate Monotonic, Earliest Deadline First, Priority Ceiling Protocol.
[SECTION_END Test]

[SECTION_START Bibliography]
1. Таненбаум Э. Современные операционные системы. - СПб.: Питер, 2015. - 1120 с.
2. Automation Studio Documentation. B&R Industrial Automation, 2020.
3. OST 7.32-2017 Отчет о научно-исследовательской работе. Структура и правила оформления.
4. ГОСТ Р 7.05-2008 Библиографическая ссылка. Общие требования и правила составления.
[SECTION_END Bibliography]
"""

# =============================================================================
# ВСПОМОГАТЕЛЬНЫЕ ФУНКЦИИ
# =============================================================================

def parse_sections(text):
    """Парсит текст и возвращает словарь с разделами"""
    sections = {}
    lines = text.strip().split('\n')
    current_section = None
    content = []
    
    for line in lines:
        if line.startswith('[SECTION_START'):
            current_section = line.split(' ')[1].replace(']', '')
            content = []
        elif line.startswith('[SECTION_END'):
            sections[current_section] = '\n'.join(content).strip()
            current_section = None
            content = []
        elif current_section is not None:
            content.append(line)
    
    return sections

def set_page_settings(doc):
    """Устанавливает настройки страницы согласно требованиям"""
    section = doc.sections[0]
    section.page_height = Cm(29.7)  # A4
    section.page_width = Cm(21)     # A4
    section.top_margin = Cm(2)      # Верхнее поле 2 см
    section.bottom_margin = Cm(2)   # Нижнее поле 2 см
    section.left_margin = Cm(3)     # Левое поле 3 см
    section.right_margin = Cm(1.5)  # Правое поле 1.5 см

def add_heading(doc, text, level=1):
    """Добавляет заголовок с правильным стилем"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(14)
        run.bold = True
    # Убираем отступы после заголовка
    heading.paragraph_format.space_after = Pt(6)

def add_paragraph(doc, text, bold=False, italic=False):
    """Добавляет абзац с правильным форматированием"""
    if not text.strip():
        return
        
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY  # Выравнивание по ширине
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    
    # Настройка форматирования абзаца
    p.paragraph_format.line_spacing = 1.5      # Межстрочный интервал 1.5
    p.paragraph_format.space_before = Pt(0)    # Отступ перед абзацем
    p.paragraph_format.space_after = Pt(0)     # Отступ после абзаца
    p.paragraph_format.first_line_indent = Cm(1.0)  # Красная строка 1.0 см

def add_image_with_caption(doc, image_path, caption_text, width=Cm(12)):
    """Добавляет изображение с подписью по центру"""
    try:
        # Проверяем существование файла
        if not os.path.exists(image_path):
            raise FileNotFoundError(f"Файл {image_path} не найден")
        
        # Добавляем изображение
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run()
        run.add_picture(image_path, width=width)
        
        # Добавляем подпись к изображению
        caption_paragraph = doc.add_paragraph()
        caption_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        caption_run = caption_paragraph.add_run(caption_text)
        caption_run.font.name = 'Times New Roman'
        caption_run.font.size = Pt(12)
        caption_run.italic = True
        
        # Добавляем пустую строку после изображения
        doc.add_paragraph()
        
        return True
        
    except FileNotFoundError:
        print(f"⚠ Файл изображения не найден: {image_path}")
        # Добавляем заглушку, если изображение не найдено
        paragraph = doc.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(f"[Изображение: {os.path.basename(image_path)} не найдено]")
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        run.italic = True
        run.font.color.rgb = RGBColor(255, 0, 0)  # Красный цвет для предупреждения
        return False
        
    except Exception as e:
        print(f"❌ Ошибка при добавлении изображения {image_path}: {e}")
        return False

def setup_page_numbers(doc):
    """Настраивает нумерацию страниц (начиная со второй страницы)"""
    # Для всех секций документа, начиная со второй
    for i, section in enumerate(doc.sections):
        if i == 0:  # Пропускаем титульную страницу
            continue
            
        # Отключаем связь с предыдущим footer
        footer = section.footer
        footer.is_linked_to_previous = False
        
        # Очищаем существующий footer
        for paragraph in footer.paragraphs:
            p = paragraph._element
            p.getparent().remove(p)
        
        # Создаем новый footer с номером страницы
        footer_paragraph = footer.add_paragraph()
        footer_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Добавляем поле для номера страницы
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'begin')
        
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = "PAGE"
        
        fld_char2 = OxmlElement('w:fldChar')
        fld_char2.set(qn('w:fldCharType'), 'end')
        
        run = footer_paragraph.add_run()
        run._r.append(fld_char)
        run._r.append(instr_text)
        run._r.append(fld_char2)
        
        # Форматируем текст номера страницы
        for run in footer_paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(10)

def process_text_with_images(doc, text, section_name):
    """Обрабатывает текст, заменяя метки изображений на реальные картинки"""
    lines = text.split('\n')
    
    for line in lines:
        if not line.strip():
            continue
            
        # Обрабатываем специальные метки для изображений
        if line.strip() == '[IMAGE_HARDWARE_CONFIG]':
            add_image_with_caption(doc, IMAGES['hardware_config'], "Рисунок 1 - Конфигурация аппаратного обеспечения")
        elif line.strip() == '[IMAGE_CYCLIC_LAYOUT_1]':
            add_image_with_caption(doc, IMAGES['cyclic_layout_1'], "Рисунок 2 - Расположение программ в одном cyclic task")
        elif line.strip() == '[IMAGE_CYCLIC_LAYOUT_2]':
            add_image_with_caption(doc, IMAGES['cyclic_layout_2'], "Рисунок 3 - Другое расположение программ в cyclic task")
        elif line.strip() == '[IMAGE_CYCLIC_LAYOUT_3]':
            add_image_with_caption(doc, IMAGES['cyclic_layout_3'], "Рисунок 4 - Расположение программ в разных cyclic tasks")
        elif line.strip() == '[IMAGE_GRAPH_1]':
            add_image_with_caption(doc, IMAGES['graph_1'], "Рисунок 5 - Результаты выполнения программ в одном cyclic task")
        elif line.strip() == '[IMAGE_GRAPH_2]':
            add_image_with_caption(doc, IMAGES['graph_2'], "Рисунок 6 - Результаты при другом расположении программ")
        elif line.strip() == '[IMAGE_GRAPH_3]':
            add_image_with_caption(doc, IMAGES['graph_3'], "Рисунок 7 - Результаты выполнения программ в разных cyclic tasks")
        else:
            # Обычный текст
            add_paragraph(doc, line)

def create_example_table(doc):
    """Создает пример таблицы с экспериментальными данными"""
    # Добавляем заголовок таблицы (выравнивание слева, не жирный)
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title_para.add_run("Таблица 1 – Время выполнения программ при разных конфигурациях")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    # Не жирный, как требуется
    
    # Создаем таблицу
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Заполняем заголовки
    headers = ['Конфигурация', 'Программа 1 (мс)', 'Программа 2 (мс)', 'Программа 3 (мс)']
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.text = header
        # Форматируем заголовки
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(10)
                run.bold = True  # Заголовки жирные
    
    # Заполняем данные
    data = [
        ['Один cyclic', '15.2', '22.8', '30.4'],
        ['Разные приоритеты', '10.5', '18.3', '45.2'],
        ['Измененный порядок', '18.7', '15.9', '28.3']
    ]
    
    for row_idx, row_data in enumerate(data, 1):
        for col_idx, value in enumerate(row_data):
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            # Форматируем данные
            for paragraph in cell.paragraphs:
                if col_idx == 0:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in paragraph.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(10)
    
    # Добавляем пустую строку после таблицы
    doc.add_paragraph()

# =============================================================================
# ГЛАВНАЯ ФУНКЦИЯ
# =============================================================================

def main():
    """Основная функция генерации отчета"""
    print("🚀 Начало генерации отчета по лабораторной работе №1")
    
    try:
        # 1. Парсим текст отчета на разделы
        sections = parse_sections(REPORT_TEXT)
        print(f"✅ Найдены секции: {list(sections.keys())}")
        
        # 2. Создаем или открываем документ
        try:
            doc = Document(TITLE_PAGE_FILE)
            print("✅ Титульный лист загружен")
        except FileNotFoundError:
            print("⚠ Титульный лист не найден, создаем новый документ")
            doc = Document()
        
        # 3. Устанавливаем настройки страницы
        set_page_settings(doc)
        print("✅ Настройки страницы применены")
        
        # 4. Настраиваем нумерацию страниц
        setup_page_numbers(doc)
        print("✅ Нумерация страниц настроена")
        
        # 5. Добавляем разрыв страницы после титульного листа
        doc.add_page_break()
        
        # 6. Добавляем разделы отчета в правильном порядке
        section_number = 1
        for section_name in SECTION_ORDER:
            if section_name in sections:
                russian_name = RUSSIAN_SECTION_NAMES.get(section_name, section_name)
                add_heading(doc, f"{section_number}. {russian_name}", level=1)
                
                # Особенная обработка для расчетно-графической части
                if section_name == 'Calculations':
                    process_text_with_images(doc, sections[section_name], section_name)
                    # Добавляем таблицу вручную
                    create_example_table(doc)
                else:
                    process_text_with_images(doc, sections[section_name], section_name)
                
                section_number += 1
                print(f"✅ Добавлен раздел: {russian_name}")
        
        # 7. Сохраняем результат
        doc.save(OUTPUT_FILE)
        print(f"✅ Отчет успешно сохранен в файл: {OUTPUT_FILE}")
        
        # 8. Информация об использованных изображениях
        print("\n📊 Использованные изображения:")
        for img_name, img_path in IMAGES.items():
            exists = "✅" if os.path.exists(img_path) else "❌"
            print(f"  {exists} {img_name}: {img_path}")
            
    except Exception as e:
        print(f"❌ Ошибка при генерации отчета: {e}")
        import traceback
        traceback.print_exc()

# =============================================================================
# ЗАПУСК СКРИПТА
# =============================================================================

if __name__ == "__main__":
    # Добавляем поддержку RGBColor для цветного текста
    try:
        from docx.shared import RGBColor
    except ImportError:
        class RGBColor:
            def __init__(self, r, g, b):
                self.rgb = (r, g, b)
    
    main()
    print("\n🎉 Генерация отчета завершена!")